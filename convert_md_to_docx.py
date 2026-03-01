# Copyright (c) Opendatalab. All rights reserved.
import subprocess
import os
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def set_docx_fonts(docx_path):
    """
    Set fonts in DOCX file: English body text to Times New Roman, Chinese text to SimSun.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        logger.error("python-docx is not installed. Please install it using 'pip install python-docx'.")
        return False

    try:
        # Open the DOCX file
        doc = Document(docx_path)

        # Set font for paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Set font for English text
                run.font.name = 'Times New Roman'
                # Set font for Chinese text
                r = run._element
                rPr = r.get_or_add_rPr()
                lang = OxmlElement('w:lang')
                lang.set(qn('w:val'), 'zh-CN')
                rPr.append(lang)
                # Set Chinese font
                font = OxmlElement('w:rFonts')
                font.set(qn('w:eastAsia'), 'SimSun')
                font.set(qn('w:ascii'), 'Times New Roman')
                font.set(qn('w:hAnsi'), 'Times New Roman')
                rPr.append(font)

        # Set font for tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # Set font for English text
                            run.font.name = 'Times New Roman'
                            # Set font for Chinese text
                            r = run._element
                            rPr = r.get_or_add_rPr()
                            lang = OxmlElement('w:lang')
                            lang.set(qn('w:val'), 'zh-CN')
                            rPr.append(lang)
                            # Set Chinese font
                            font = OxmlElement('w:rFonts')
                            font.set(qn('w:eastAsia'), 'SimSun')
                            font.set(qn('w:ascii'), 'Times New Roman')
                            font.set(qn('w:hAnsi'), 'Times New Roman')
                            rPr.append(font)

        # Save the modified DOCX file
        doc.save(docx_path)
        logger.info(f"Successfully set fonts in {docx_path}")
        return True
    except Exception as e:
        logger.exception(f"Error setting fonts in {docx_path}: {e}")
        return False


def html_table_to_pipe_table(html_table: str) -> str:
    """
    将 HTML 表格转换为标准的 pipe table 格式。
    """
    import re

    # 提取所有行
    rows = re.findall(r'<tr[^>]*>(.*?)</tr>', html_table, re.DOTALL | re.IGNORECASE)
    if not rows:
        return html_table

    # 提取单元格的函数
    def extract_cells(row_html):
        cells = []
        # 匹配 th 和 td 标签
        cell_matches = re.findall(r'<(th|td)[^>]*>(.*?)</\1>', row_html, re.DOTALL | re.IGNORECASE)
        for _, cell_content in cell_matches:
            # 去除 HTML 标签，只保留文本
            text = re.sub(r'<[^>]*>', '', cell_content)
            # 去除空白字符
            text = text.strip()
            # 替换 HTML 实体
            text = text.replace('&gt;', '>').replace('&lt;', '<').replace('&amp;', '&')
            cells.append(text)
        return cells

    # 处理表头
    header_cells = extract_cells(rows[0])
    if not header_cells:
        return html_table

    # 处理数据行
    data_rows = []
    for row in rows[1:]:
        data_cells = extract_cells(row)
        if data_cells:
            data_rows.append(data_cells)

    # 计算每列的最大宽度
    max_widths = []
    all_rows = [header_cells] + data_rows
    for i in range(len(header_cells)):
        max_width = max(len(row[i]) for row in all_rows if i < len(row))
        max_widths.append(max_width)

    # 构建 pipe table
    pipe_table = []

    # 添加表头
    header_line = '| '
    for i, cell in enumerate(header_cells):
        padding = ' ' * (max_widths[i] - len(cell))
        header_line += cell + padding + ' | '
    pipe_table.append(header_line.rstrip())

    # 添加分隔线
    separator_line = '| '
    for width in max_widths:
        separator_line += '-' * width + ' | '
    pipe_table.append(separator_line.rstrip())

    # 添加数据行
    for row in data_rows:
        data_line = '| '
        for i, cell in enumerate(row):
            if i < len(max_widths):
                padding = ' ' * (max_widths[i] - len(cell))
                data_line += cell + padding + ' | '
        pipe_table.append(data_line.rstrip())

    return '\n'.join(pipe_table)


def ensure_table_surrounded_by_blank_lines(md_content: str) -> str:
    """
    确保每个 <table>...</table> 块前后都有一个空行（即前后是 \n\n）。
    处理跨多行的表格，兼容任意空白字符。
    """
    import re
    # 正则匹配整个 <table> ... </table> 块（非贪婪，支持跨行）
    def add_blank_lines(match):
        table_block = match.group(0)
        # 确保前面有 \n\n（但不在字符串开头强行加两个换行）
        if not table_block.startswith('\n\n'):
            # 在前面加两个换行，但要替换掉可能已有的部分换行
            table_block = '\n\n' + table_block.lstrip('\n')
        # 确保后面有 \n\n（但不在字符串末尾强行加两个换行）
        if not table_block.endswith('\n\n'):
            table_block = table_block.rstrip('\n') + '\n\n'
        return table_block

    # 使用 re.DOTALL 使 . 匹配换行符，re.IGNORECASE 忽略大小写（如 <TABLE>）
    # 注意：只匹配完整的 <table>...</table> 对
    pattern = r'<table\b[^>]*>.*?</table>'
    result = re.sub(pattern, add_blank_lines, md_content, flags=re.DOTALL | re.IGNORECASE)

    # 清理可能产生的过多连续空行（可选，保持整洁）
    result = re.sub(r'\n{3,}', '\n\n', result)

    return result.strip() + '\n'  # 保证文件以单个换行结尾


def fix_table_headers(md_file):
    """
    Fix table headers in markdown file by replacing <td> with <th> in the first row of each table.
    Also ensure tables are surrounded by blank lines and convert HTML tables to pipe tables.
    """
    try:
        import re
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # 确保表格前后有空白行
        content = ensure_table_surrounded_by_blank_lines(content)

        # 匹配所有表格并转换为 pipe table
        def convert_to_pipe_table(match):
            html_table = match.group(0)
            pipe_table = html_table_to_pipe_table(html_table)
            return pipe_table

        # 匹配所有表格并处理
        content_fixed = re.sub(r'<table\b[^>]*>.*?</table>', convert_to_pipe_table, content, flags=re.DOTALL | re.IGNORECASE)

        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(content_fixed)
        logger.info(f"Fixed table headers in {md_file}")
        return True
    except Exception as e:
        logger.exception(f"Error fixing table headers in {md_file}: {e}")
        return False


def convert_md_to_docx():
    """
    Convert all Markdown files under aims directory to DOCX format using pandoc,
    ensuring titles, tables, formulas, and images are properly handled.
    Then set fonts: English body text to Times New Roman, Chinese text to SimSun.
    """
    aims_dir = Path("./aims")

    if not aims_dir.exists():
        logger.error(f"Directory {aims_dir} does not exist.")
        return

    # Find all markdown files under aims directory
    md_files = list(aims_dir.rglob("*.md"))

    if not md_files:
        logger.info(f"No markdown files found under {aims_dir}.")
        return

    logger.info(f"Found {len(md_files)} markdown files to convert.")

    # Check if pandoc is installed
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
    except (subprocess.SubprocessError, FileNotFoundError):
        logger.error("pandoc is not installed. Please install pandoc first.")
        logger.info("Visit https://pandoc.org/installing.html for installation instructions.")
        return

    # Convert each markdown file to docx
    for md_file in md_files:
        try:
            docx_file = md_file.with_suffix(".docx")
            logger.info(f"Converting {md_file} to {docx_file}...")

            # Fix table headers before conversion
            logger.info(f"Fixing table headers in {md_file}...")
            fix_table_headers(str(md_file))

            # Change to the directory of the markdown file to ensure relative paths work
            md_dir = md_file.parent

            # Run pandoc command with optimal options for table, image, formula, and heading conversion
            # Based on Pandoc's documentation, gfm format has the best table support
            # Using basic markdown format to ensure compatibility with all table types
            result = subprocess.run(
                [
                    "pandoc",
                    md_file.name,
                    "-o", str(docx_file.relative_to(md_dir)),
                    "--extract-media=media",   # Extract images to media directory
                    "--standalone",       # Create a standalone document
                    #"--from=markdown+pipe_tables+tex_math_dollars",  # Use basic markdown with pipe tables and LaTeX formulas
                    "--from=markdown+raw_html+grid_tables+pipe_tables+simple_tables+multiline_tables+table_captions+tex_math_dollars",
                    "--to=docx",          # Explicitly set output format to docx
                    "--mathml",           # Convert LaTeX formulas to MathML for better docx compatibility
                    "--table-of-contents", # Generate table of contents
                    "--number-sections" ,   # Number sections for better navigation
                    "--toc-depth=3"
                ],
                cwd=str(md_dir),
                capture_output=True,
                text=True
            )

            if result.returncode == 0:
                logger.info(f"Successfully converted {md_file} to {docx_file}.")
                # Check if media directory was created and has files
                media_dir = md_dir / "media"
                if media_dir.exists() and any(media_dir.iterdir()):
                    logger.info(f"Media files (images) extracted to: {media_dir}")
                else:
                    logger.info("No media files found or extracted.")

                # Set fonts in the converted DOCX file
                logger.info(f"Setting fonts in {docx_file}...")
                set_docx_fonts(str(docx_file))
            else:
                logger.error(f"Failed to convert {md_file}: {result.stderr}")

        except Exception as e:
            logger.exception(f"Error converting {md_file}: {e}")

    logger.info("Conversion process completed.")


if __name__ == "__main__":
    convert_md_to_docx()
